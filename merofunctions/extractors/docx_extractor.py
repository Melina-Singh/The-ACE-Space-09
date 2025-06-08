import logging
from docx import Document
from io import BytesIO
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_text(text):
    """Remove extra whitespace and special characters, keep Unicode."""
    try:
        text = re.sub(r'\s+', ' ', text.strip())
        text = re.sub(r'[^\w\s.,\-–—:;!?()\'\"%€$£]', '', text, flags=re.UNICODE)
        return text
    except Exception as e:
        logger.error(f"Text preprocessing failed: {str(e)}")
        raise

def extract_docx(blob_data, chunk_size=10000):
    """
    Extract and preprocess text from a DOCX blob, with chunking for large files.
    Args:
        blob_data (bytes): DOCX file content.
        chunk_size (int): Number of characters per chunk.
    Returns:
        dict: List of text chunks and tables.
    """
    try:
        logger.info("Starting DOCX extraction")
        doc = Document(BytesIO(blob_data))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = clean_text(" ".join(paragraphs))
        text_chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        logger.info(f"Extracted {len(paragraphs)} paragraphs, {len(text_chunks)} text chunks, and {len(tables)} tables from DOCX")
        return {"text_chunks": text_chunks, "tables": tables, "type": "docx"}
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise